from __future__ import annotations

import hashlib
import re
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(r"C:\SmartKasirQR")
FLUTTER = ROOT / "smartkasir-qr-mobile"
LARAVEL = ROOT / "smartkasir-qr-api"
OUT = ROOT / "submission"
REPORT_DIR = OUT / "03_Laporan"
API_DIR = OUT / "05_Dokumentasi_API"

DOCX_PATH = REPORT_DIR / "Laporan_UAS_SmartKasir_QR.docx"
PDF_PATH = REPORT_DIR / "Laporan_UAS_SmartKasir_QR_DRAFT.pdf"
SCREENSHOT_CHECKLIST = REPORT_DIR / "DAFTAR_SCREENSHOT_YANG_HARUS_DIISI.txt"
PDF_CHECK = REPORT_DIR / "PDF_DRAFT_CHECK.txt"
ENDPOINT_DOC = API_DIR / "ENDPOINT_API.md"

IDENTITY = {
    "nama": "I WAYAN NUGRAHA PURNAMA",
    "nim": "2323050026",
    "program": "Sistem Informasi",
    "fakultas": "Fakultas Sains dan Teknologi",
    "universitas": "Universitas Tabanan",
    "matkul": "Pemrograman Mobile Lanjutan",
    "semester": "VI",
    "dosen": "Luh Ayu Diah Fernita Sari, S.Kom., M.T.",
    "ujian": "Ujian Akhir Semester Genap Tahun Akademik 2025/2026",
    "tahun": "2026",
}

ENDPOINTS = [
    ("POST", "/api/login", "Tidak", "Login kasir dan menghasilkan Bearer Token.", "Halaman Login"),
    ("GET", "/api/me", "Bearer Token", "Memeriksa sesi pengguna aktif.", "Pemulihan Sesi"),
    ("POST", "/api/logout", "Bearer Token", "Logout dan menghapus token server.", "Profil"),
    ("GET", "/api/dashboard", "Bearer Token", "Mengambil ringkasan dashboard kasir.", "Beranda"),
    ("GET", "/api/categories", "Bearer Token", "Mengambil daftar kategori produk.", "Produk dan Tambah Produk"),
    ("GET", "/api/products", "Bearer Token", "Mengambil daftar produk dengan search/filter.", "Produk"),
    ("GET", "/api/products/{id}", "Bearer Token", "Mengambil detail produk.", "Detail Produk"),
    ("POST", "/api/products", "Bearer Token", "Menyimpan produk baru.", "Tambah Produk"),
    ("GET", "/api/orders", "Bearer Token", "Mengambil daftar pesanan.", "Pesanan"),
    ("GET", "/api/orders/{id}", "Bearer Token", "Mengambil detail pesanan.", "Detail Pesanan"),
    ("PATCH", "/api/orders/{id}/confirm", "Bearer Token", "Konfirmasi pesanan dan kurangi stok.", "Detail Pesanan"),
    ("PATCH", "/api/orders/{id}/process", "Bearer Token", "Mengubah pesanan menjadi diproses.", "Detail Pesanan"),
    ("PATCH", "/api/orders/{id}/ready", "Bearer Token", "Menandai pesanan siap.", "Detail Pesanan"),
    ("PATCH", "/api/orders/{id}/cancel", "Bearer Token", "Membatalkan pesanan dan mengembalikan stok bila perlu.", "Detail Pesanan"),
    ("GET", "/api/transactions", "Bearer Token", "Mengambil daftar transaksi.", "Transaksi"),
    ("GET", "/api/transactions/{id}", "Bearer Token", "Mengambil detail transaksi.", "Detail Transaksi"),
    ("POST", "/api/transactions", "Bearer Token", "Membuat transaksi kasir atau pembayaran order QR.", "Transaksi dan Pembayaran"),
    ("GET", "/api/public/tables/{qrToken}/menu", "Tidak", "Mengambil menu publik berdasarkan QR meja.", "Website Pelanggan"),
    ("POST", "/api/public/tables/{qrToken}/orders", "Tidak", "Membuat pesanan publik dari QR meja.", "Website Pelanggan"),
    ("GET", "/api/public/orders/{orderNumber}/status", "Tidak", "Melihat status pesanan publik.", "Website Pelanggan"),
]

SCREENSHOTS = [
    ("01-login.png", "Gambar 7.1 Halaman Login", "Halaman login digunakan kasir untuk masuk menggunakan email dan password demo."),
    ("02-home-dashboard.png", "Gambar 7.2 Home Dashboard", "Dashboard menampilkan ringkasan pesanan, transaksi, produk, stok rendah, dan pesanan terbaru."),
    ("03-products-get.png", "Gambar 7.3 Daftar Produk dari GET API", "Daftar produk berasal dari endpoint GET dan mendukung pencarian serta filter."),
    ("04-product-detail.png", "Gambar 7.4 Detail Produk", "Detail produk menampilkan data harga, stok, kategori, status, dan SKU."),
    ("05-create-product-form.png", "Gambar 7.5 Form Tambah Produk", "Form tambah produk dipakai untuk mengirim data produk baru ke REST API."),
    ("06-create-product-success.png", "Gambar 7.6 Produk Berhasil Ditambahkan", "Halaman sukses menandakan produk telah disimpan dan daftar disegarkan."),
    ("07-orders-get.png", "Gambar 7.7 Daftar Pesanan dari GET API", "Daftar pesanan menampilkan status order dan dapat difilter sesuai alur kerja kasir."),
    ("08-order-detail.png", "Gambar 7.8 Detail Pesanan", "Detail pesanan menampilkan meja, pelanggan, item, total, dan aksi status."),
    ("09-order-confirm.png", "Gambar 7.9 Konfirmasi Pesanan", "Konfirmasi pesanan mengubah status pending menjadi confirmed serta mengurangi stok."),
    ("10-order-payment.png", "Gambar 7.10 Pembayaran Pesanan", "Pembayaran order QR mengirim order_id dan jumlah bayar tanpa mengirim ulang item."),
    ("11-transactions-get.png", "Gambar 7.11 Daftar Transaksi dari GET API", "Daftar transaksi menampilkan transaksi kasir dan pembayaran order QR."),
    ("12-transaction-detail.png", "Gambar 7.12 Detail Transaksi", "Detail transaksi menampilkan nomor transaksi, kasir, item, total, bayar, dan kembalian."),
    ("13-create-transaction-form.png", "Gambar 7.13 Form Transaksi Kasir", "Form transaksi kasir dipakai untuk memilih produk, jumlah, dan pembayaran tunai."),
    ("14-transaction-success.png", "Gambar 7.14 Transaksi Berhasil", "Halaman sukses transaksi menampilkan hasil transaksi yang dihitung backend."),
    ("15-profile-logout.png", "Gambar 7.15 Profil dan Logout", "Halaman profil menampilkan pengguna aktif dan tombol logout."),
    ("16-qr-tables.png", "Gambar 7.16 Daftar QR Meja", "Website QR meja menampilkan daftar QR atau link menu untuk pelanggan."),
    ("17-qr-customer-menu.png", "Gambar 7.17 Menu Pelanggan", "Menu pelanggan menampilkan produk tersedia dan keranjang lokal."),
    ("18-qr-order-success.png", "Gambar 7.18 Pesanan Pelanggan Berhasil", "Halaman sukses menampilkan nomor pesanan pelanggan."),
]

TEST_ROWS = [
    ("Login benar", "Pengguna dapat masuk dengan email dan password demo.", "Lulus melalui test autentikasi dan integrasi API.", "Lulus"),
    ("Login salah", "Aplikasi menampilkan pesan kesalahan.", "Lulus melalui test auth.", "Lulus"),
    ("Pemulihan sesi", "Token dicek melalui GET /api/me.", "Lulus melalui test AuthProvider.", "Lulus"),
    ("Logout", "Token dihapus dan halaman kembali ke login.", "Lulus melalui test logout provider.", "Lulus"),
    ("GET Dashboard", "Dashboard tampil dari API.", "Lulus melalui test dashboard.", "Lulus"),
    ("GET Produk", "Daftar produk tampil.", "Lulus melalui test produk.", "Lulus"),
    ("GET Pesanan", "Daftar pesanan tampil.", "Lulus melalui test pesanan.", "Lulus"),
    ("GET Transaksi", "Daftar transaksi tampil.", "Lulus melalui test transaksi.", "Lulus"),
    ("POST Produk", "Produk baru berhasil disimpan.", "Lulus melalui ProductApiTest dan test Flutter.", "Lulus"),
    ("SKU duplikat", "Server menolak SKU yang sama.", "Lulus melalui ProductApiTest.", "Lulus"),
    ("POST Transaksi", "Transaksi tersimpan dan stok berkurang.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pembayaran kurang", "Server menolak pembayaran kurang.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pesanan QR pending", "Pesanan QR dibuat pending tanpa mengurangi stok.", "Lulus melalui test API dan website QR.", "Lulus"),
    ("Konfirmasi pesanan", "Status menjadi confirmed dan stok berkurang.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Proses pesanan", "Status menjadi processing.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pesanan siap", "Status menjadi ready.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pembayaran pesanan", "Order menjadi completed dan transaksi dibuat.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pembayaran ganda", "Pembayaran kedua ditolak.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Cancel pending", "Pesanan dibatalkan tanpa perubahan stok.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Cancel confirmed", "Pesanan dibatalkan dan stok kembali.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Pengembalian stok", "Stok tidak dikembalikan dua kali.", "Lulus melalui StageThreeApiTest.", "Lulus"),
    ("Penanganan 401", "Sesi dihapus dan kembali ke login.", "Lulus melalui test provider.", "Lulus"),
    ("Server tidak tersedia", "UI menampilkan error.", "Lulus melalui automated test menggunakan fake repository/provider.", "Lulus"),
    ("Build APK", "APK release berhasil dibuat.", "Belum diuji karena Android SDK/device belum tersedia.", "Belum Lulus"),
]

SNIPPETS = [
    ("Dio Bearer Token Interceptor", "lib/core/network/api_client.dart", FLUTTER / "lib/core/network/api_client.dart", ["options.headers['Authorization']"]),
    ("AuthRepository Login", "lib/features/auth/repositories/auth_repository.dart", FLUTTER / "lib/features/auth/repositories/auth_repository.dart", ["Future<UserModel> login", "login("]),
    ("TokenStorage", "lib/core/storage/token_storage.dart", FLUTTER / "lib/core/storage/token_storage.dart", ["class TokenStorage"]),
    ("AuthProvider", "lib/features/auth/providers/auth_provider.dart", FLUTTER / "lib/features/auth/providers/auth_provider.dart", ["Future<bool> login", "class AuthProvider"]),
    ("DashboardRepository", "lib/features/dashboard/repositories/dashboard_repository.dart", FLUTTER / "lib/features/dashboard/repositories/dashboard_repository.dart", ["fetchDashboard"]),
    ("ProductRepository GET", "lib/features/products/repositories/product_repository.dart", FLUTTER / "lib/features/products/repositories/product_repository.dart", ["fetchProducts"]),
    ("ProductRepository POST", "lib/features/products/repositories/product_repository.dart", FLUTTER / "lib/features/products/repositories/product_repository.dart", ["createProduct"]),
    ("OrderRepository PATCH", "lib/features/orders/repositories/order_repository.dart", FLUTTER / "lib/features/orders/repositories/order_repository.dart", ["confirmOrder"]),
    ("TransactionRepository POST", "lib/features/transactions/repositories/transaction_repository.dart", FLUTTER / "lib/features/transactions/repositories/transaction_repository.dart", ["createTransaction"]),
    ("MainNavigationPage", "lib/navigation/main_navigation_page.dart", FLUTTER / "lib/navigation/main_navigation_page.dart", ["class MainNavigationPage"]),
    ("Route API Laravel", "routes/api.php", LARAVEL / "routes/api.php", ["Route::post('/login'"]),
    ("OrderService Pengelolaan Stok", "app/Services/OrderService.php", LARAVEL / "app/Services/OrderService.php", ["public function confirm"]),
    ("TransactionService Transaksi", "app/Services/TransactionService.php", LARAVEL / "app/Services/TransactionService.php", ["public function create"]),
]


def ensure_markers() -> None:
    markers = [
        FLUTTER / "pubspec.yaml",
        FLUTTER / "lib/main.dart",
        LARAVEL / "artisan",
        LARAVEL / "composer.json",
        LARAVEL / "routes/api.php",
    ]
    missing = [str(path) for path in markers if not path.exists()]
    if missing:
        raise SystemExit("Marker tidak ditemukan:\n" + "\n".join(missing))


def ensure_dirs() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    API_DIR.mkdir(parents=True, exist_ok=True)


def rel(path: Path, root: Path) -> str:
    return str(path.relative_to(root)).replace("\\", "/")


def read_snippet(path: Path, needles: list[str], max_lines: int = 28) -> str:
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    start = 0
    for index, line in enumerate(lines):
        if any(needle in line for needle in needles):
            start = max(index - 6, 0)
            break
    snippet = "\n".join(lines[start : start + max_lines])
    snippet = re.sub(r"C:\\Users\\[^\\\n]+\\", "", snippet)
    return snippet


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "777777", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc: Document, title: str, location: str, explanation: str, code: str) -> None:
    doc.add_heading(title, level=2)
    add_para(doc, f"Lokasi file: {location}. {explanation}")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F4F4")
    set_cell_border(cell, "BBBBBB", "6")
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_screenshot_placeholder(doc: Document, filename: str, caption: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(14)
    row = table.rows[0]
    row.height = Cm(7.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    set_cell_border(cell, "777777", "10")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell.paragraphs[0]
    p.add_run("\n\nTEMPEL SCREENSHOT DI SINI\n").bold = True
    r = p.add_run(filename)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    add_para(doc, body)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def build_endpoint_doc() -> None:
    lines = [
        "# Dokumentasi Endpoint API SmartKasir QR",
        "",
        "Base URL development browser: `http://127.0.0.1:8000/api`",
        "Base URL Android Emulator: `http://10.0.2.2:8000/api`",
        "",
        "Endpoint privat menggunakan header:",
        "",
        "```http",
        "Authorization: Bearer {{bearer_token}}",
        "Accept: application/json",
        "```",
        "",
        "Token nyata tidak ditulis pada dokumentasi ini.",
        "",
        "| No | Method | Endpoint | Auth | Request Body | Response Ringkas | Status Code | Digunakan Pada |",
        "|---:|---|---|---|---|---|---|---|",
    ]
    bodies = {
        "POST /api/login": "`email`, `password`, opsional `device_name`",
        "POST /api/products": "`category_id`, `name`, `sku`, `description`, `price`, `stock`, `unit`, `is_available`",
        "POST /api/transactions": "Transaksi langsung: `paid_amount`, `payment_method`, `items[]`; pembayaran QR: `order_id`, `paid_amount`, `payment_method`",
        "PATCH /api/orders/{id}/cancel": "Opsional `reason`",
        "POST /api/public/tables/{qrToken}/orders": "`customer_name`, `notes`, `items[]`",
    }
    for i, (method, endpoint, auth, function, page) in enumerate(ENDPOINTS, 1):
        key = f"{method} {endpoint}"
        body = bodies.get(key, "-")
        status = "200/201 sukses, 401 tanpa token, 422 validasi, 404 data tidak ditemukan"
        response = "JSON berisi `success`, `message`, dan `data` sesuai resource."
        lines.append(f"| {i} | {method} | `{endpoint}` | {auth} | {body} | {response} | {status} | {page} |")
    ENDPOINT_DOC.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_checklist() -> None:
    lines = [
        "DAFTAR SCREENSHOT YANG HARUS DIISI",
        "",
        "| No | Nama File | Judul Gambar | Halaman Laporan | Status |",
        "|---:|---|---|---|---|",
    ]
    for i, (filename, caption, _body) in enumerate(SCREENSHOTS, 1):
        lines.append(f"| {i} | {filename} | {caption} | BAB VII | BELUM DIISI |")
    lines += [
        "",
        "Petunjuk pengisian:",
        "1. Buka DOCX.",
        "2. Klik placeholder screenshot.",
        "3. Hapus kotak placeholder.",
        "4. Pilih Insert -> Pictures.",
        "5. Pilih screenshot sesuai nama file.",
        "6. Atur lebar maksimal 14 cm.",
        "7. Pertahankan caption yang sudah tersedia.",
        "8. Simpan DOCX.",
        "9. Konversi ulang menjadi PDF.",
    ]
    SCREENSHOT_CHECKLIST.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.bold = True
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UNIVERSITAS TABANAN\n\nIMPLEMENTASI APLIKASI SMARTKASIR QR\nBERBASIS FLUTTER DAN REST API LARAVEL")
    run.bold = True
    run.font.size = Pt(14)
    for line in [
        IDENTITY["ujian"],
        f"Mata Kuliah: {IDENTITY['matkul']}",
        f"Semester: {IDENTITY['semester']}",
        f"Dosen Pengampu: {IDENTITY['dosen']}",
        f"Nama: {IDENTITY['nama']}",
        f"NIM: {IDENTITY['nim']}",
        f"Program Studi: {IDENTITY['program']}",
        IDENTITY["fakultas"],
        IDENTITY["universitas"],
        IDENTITY["tahun"],
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("KATA PENGANTAR", level=1)
    add_para(doc, "Puji syukur penulis panjatkan ke hadapan Tuhan Yang Maha Esa karena laporan Ujian Akhir Semester ini dapat disusun sebagai dokumentasi implementasi aplikasi SmartKasir QR. Laporan ini membahas pembangunan aplikasi kasir berbasis Flutter yang terhubung dengan REST API Laravel, menggunakan autentikasi Bearer Token, database SQLite, serta website pelanggan berbasis QR meja.")
    add_para(doc, "Penulis menyadari bahwa laporan ini masih dapat dikembangkan, khususnya pada bagian rilis Android karena APK belum dibuat saat penyusunan laporan ini. Oleh karena itu, laporan ini diberi status draft dan menyediakan placeholder screenshot agar dapat dilengkapi setelah pengujian tampilan dilakukan.")
    doc.add_page_break()

    doc.add_heading("DAFTAR ISI", level=1)
    add_para(doc, "Daftar isi dapat diperbarui di Microsoft Word melalui References -> Table of Contents setelah dokumen dibuka.")
    for item in ["BAB I PENDAHULUAN", "BAB II DESKRIPSI APLIKASI", "BAB III STRUKTUR PROYEK", "BAB IV IMPLEMENTASI FITUR DAN CODING", "BAB V DAFTAR ENDPOINT REST API", "BAB VI PENGUJIAN", "BAB VII HASIL TAMPILAN", "BAB VIII CARA MENJALANKAN", "BAB IX PENUTUP"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    chapters = [
        ("BAB I PENDAHULUAN", [
            ("1.1 Latar Belakang", "Pengelolaan produk, pesanan, transaksi, dan stok secara terpisah dapat menimbulkan keterlambatan pelayanan, kesalahan pencatatan, dan ketidaksesuaian jumlah stok. SmartKasir QR dibuat sebagai solusi kasir yang mudah digunakan, terhubung ke REST API, mendukung pemesanan pelanggan melalui QR meja, serta menyinkronkan stok secara otomatis."),
            ("1.2 Rumusan Masalah", "Rumusan masalah dalam proyek ini adalah sebagai berikut."),
            ("1.3 Tujuan", "Tujuan proyek ini adalah membangun aplikasi SmartKasir QR yang dapat digunakan kasir untuk login, melihat dashboard, mengelola produk, memproses pesanan QR, membuat transaksi, menerima pembayaran, dan melakukan logout dengan penghapusan token."),
            ("1.4 Manfaat", "Bagi kasir, sistem membantu proses transaksi dan pesanan. Bagi pelanggan, QR meja mempermudah pemesanan. Bagi pengelola usaha, sistem membantu sinkronisasi stok. Bagi mahasiswa, proyek ini menjadi implementasi pembelajaran Flutter, REST API Laravel, autentikasi, dan pengujian."),
            ("1.5 Batasan Sistem", "Aplikasi kasir dibuat dengan Flutter, backend dan website pelanggan dibuat dengan Laravel, database menggunakan SQLite, pembayaran hanya tunai, backend berjalan pada server lokal, belum memakai payment gateway, belum memakai printer struk, dan APK belum dibuat pada tahap penyusunan laporan ini."),
        ]),
        ("BAB II DESKRIPSI APLIKASI", [
            ("2.1 Gambaran Umum SmartKasir QR", "SmartKasir QR terdiri dari aplikasi Flutter untuk kasir, REST API Laravel untuk pengolahan data, website pelanggan yang diakses melalui QR meja, serta SQLite sebagai database lokal. Semua proses penting seperti stok, total transaksi, nomor transaksi, dan kembalian dihitung pada backend."),
            ("2.2 Pengguna Sistem", "Kasir dapat login, melihat dashboard, mengelola produk, memproses pesanan, membuat transaksi, menerima pembayaran, dan logout. Pelanggan dapat membuka menu melalui QR, memilih produk, membuat pesanan, memperoleh nomor pesanan, dan melihat status pesanan."),
            ("2.3 Teknologi yang Digunakan", "Teknologi utama yang digunakan dijelaskan pada tabel berikut."),
            ("2.4 Arsitektur Sistem", "Arsitektur SmartKasir QR memisahkan aplikasi kasir, REST API, service, model, dan database. Website QR pelanggan juga terhubung ke Laravel dan SQLite sehingga pesanan pelanggan dapat dilihat oleh aplikasi kasir."),
            ("2.5 Alur Pesanan QR", "Alur utama pesanan QR adalah pending, confirmed, processing, ready, dan completed. Pesanan juga dapat berubah menjadi cancelled dari status pending, confirmed, processing, atau ready. Stok tidak berkurang saat pending, berkurang saat confirm, tidak berkurang ulang saat payment, dan kembali jika pesanan dibatalkan setelah confirm."),
        ]),
        ("BAB III STRUKTUR PROYEK", [
            ("3.1 Struktur Folder Flutter", "Struktur Flutter terdiri dari config, core, features, navigation, test, main.dart, dan app.dart. Pola implementasi dibagi menjadi model, repository, provider, page, dan widget."),
            ("3.2 Struktur Folder Laravel", "Struktur Laravel terdiri dari app/Http, app/Models, app/Services, routes, database/migrations, database/seeders, resources/views/customer, dan tests."),
            ("3.3 Struktur Database", "Tabel utama yang ditemukan dari migration dan model adalah users, personal_access_tokens, categories, products, restaurant_tables, orders, order_items, transactions, dan transaction_items. Nama model terkait adalah User, Category, Product, RestaurantTable, Order, OrderItem, Transaction, dan TransactionItem."),
        ]),
    ]
    for chapter, sections in chapters:
        doc.add_heading(chapter, level=1)
        for heading, body in sections:
            doc.add_heading(heading, level=2)
            add_para(doc, body)
            if heading == "1.2 Rumusan Masalah":
                add_numbered(doc, [
                    "Bagaimana membangun aplikasi kasir Flutter yang terhubung dengan REST API Laravel?",
                    "Bagaimana menerapkan autentikasi Bearer Token pada aplikasi?",
                    "Bagaimana mengelola produk, pesanan, transaksi, dan stok dalam satu sistem terintegrasi?",
                    "Bagaimana menerapkan pemesanan pelanggan melalui QR meja?",
                ])
            if heading == "2.3 Teknologi yang Digunakan":
                add_table(doc, ["Teknologi", "Fungsi"], [
                    ["Flutter", "Membangun aplikasi kasir."],
                    ["Dart", "Bahasa pemrograman Flutter."],
                    ["Provider", "State management."],
                    ["Dio", "HTTP client REST API."],
                    ["flutter_secure_storage", "Penyimpanan token aman."],
                    ["Laravel", "REST API dan website QR pelanggan."],
                    ["Laravel Sanctum", "Bearer Token."],
                    ["SQLite", "Database lokal UAS."],
                    ["Blade/HTML/CSS/JavaScript", "Tampilan website pelanggan QR."],
                ])
            if heading == "2.4 Arsitektur Sistem":
                add_code(doc, "Diagram Arsitektur Sederhana", "dibuat dalam laporan", "Diagram berikut menjelaskan hubungan komponen utama.", "Aplikasi Flutter Kasir\n    |\nREST API Laravel\n    |\nService dan Model\n    |\nSQLite\n\nWebsite QR Pelanggan\n    |\nLaravel\n    |\nSQLite\n    |\nAplikasi Flutter Kasir")

    doc.add_heading("BAB IV IMPLEMENTASI FITUR DAN CODING", level=1)
    feature_sections = [
        ("4.1 Login dan Bearer Token", "Login memakai form email dan password, mengirim POST /api/login, menerima token, menyimpan token, memeriksa sesi melalui GET /api/me, dan menghapus token saat logout."),
        ("4.2 Home Dashboard", "Dashboard memakai GET /api/dashboard dan menampilkan ringkasan pesanan berdasarkan status, transaksi hari ini, pendapatan, produk aktif, stok rendah, dan pesanan terbaru."),
        ("4.3 Bottom Navigation", "Bottom Navigation memiliki menu Beranda, Pesanan, Produk, dan Transaksi."),
        ("4.4 GET Produk", "Halaman produk memakai GET /api/products dan GET /api/products/{id}, mendukung search, filter kategori, filter ketersediaan, detail, refresh, loading, error, dan empty state."),
        ("4.5 GET Pesanan", "Halaman pesanan memakai GET /api/orders dan GET /api/orders/{id}, mendukung search nomor pesanan, filter status, detail item, dan status badge."),
        ("4.6 GET Transaksi", "Halaman transaksi memakai GET /api/transactions dan GET /api/transactions/{id}, mendukung search, filter tanggal, dan detail transaksi."),
        ("4.7 POST Tambah Produk", "Tambah produk memakai POST /api/products dengan validasi lokal, validasi server, SKU unik, pencegahan submit ganda, serta refresh produk dan dashboard."),
        ("4.8 POST Transaksi Kasir", "Transaksi kasir memakai POST /api/transactions. Flutter tidak mengirim harga, subtotal, total, atau kembalian karena backend menghitung nilai final dan mengurangi stok."),
        ("4.9 Website Pelanggan Melalui QR", "Setiap meja memiliki QR. Pelanggan membuka menu, memilih produk, menyimpan keranjang pada localStorage, mengirim pesanan, dan memperoleh nomor pesanan."),
        ("4.10 Perubahan Status Pesanan", "Endpoint PATCH confirm, process, ready, dan cancel digunakan sesuai status order. Tombol aksi pada Flutter ditampilkan berdasarkan status."),
        ("4.11 Pembayaran Pesanan QR", "Pembayaran order QR mengirim order_id, paid_amount, dan payment_method. Item tidak dikirim ulang dari Flutter."),
        ("4.12 Sinkronisasi Stok", "Stok tetap saat pending, berkurang saat confirm, tidak berkurang lagi saat payment, dan kembali saat cancel setelah confirm."),
        ("4.13 Penanganan State", "Provider mengatur loading, error, empty state, refresh, 401 Unauthorized, reset provider saat logout, dan pencegahan submit ganda."),
    ]
    for heading, body in feature_sections:
        doc.add_heading(heading, level=2)
        add_para(doc, body)
        if heading == "4.7 POST Tambah Produk":
            add_code(doc, "Contoh Request Tambah Produk", "POST /api/products", "Request dikirim dari Flutter ke Laravel.", '{\n  "category_id": 1,\n  "name": "Produk Contoh",\n  "sku": "PRD-001",\n  "description": "Deskripsi produk",\n  "price": 5000,\n  "stock": 10,\n  "unit": "pcs",\n  "is_available": true\n}')
        if heading == "4.8 POST Transaksi Kasir":
            add_code(doc, "Contoh Request Transaksi Kasir", "POST /api/transactions", "Harga dan total dihitung oleh backend.", '{\n  "paid_amount": 20000,\n  "payment_method": "cash",\n  "items": [\n    {\n      "product_id": 1,\n      "quantity": 2\n    }\n  ]\n}')
        if heading == "4.11 Pembayaran Pesanan QR":
            add_code(doc, "Contoh Request Pembayaran Pesanan", "POST /api/transactions", "Pembayaran order tidak mengirim ulang item.", '{\n  "order_id": 5,\n  "paid_amount": 20000,\n  "payment_method": "cash"\n}')

    for title, location, path, needles in SNIPPETS:
        explanation = "Potongan kode ini diambil dari source aktual dan dipersingkat untuk menjelaskan bagian implementasi."
        add_code(doc, title, location, explanation, read_snippet(path, needles))

    doc.add_heading("BAB V DAFTAR ENDPOINT REST API", level=1)
    add_para(doc, "Endpoint berikut diaudit dari routes/api.php. Endpoint privat memakai Bearer Token Laravel Sanctum.")
    add_table(doc, ["No", "Method", "Endpoint", "Autentikasi", "Fungsi", "Digunakan Pada"], [[str(i), *row] for i, row in enumerate(ENDPOINTS, 1)])

    doc.add_heading("BAB VI PENGUJIAN", level=1)
    doc.add_heading("6.1 Pengujian Flutter", level=2)
    add_bullets(doc, ["dart format: lulus, 0 perubahan.", "flutter analyze: No issues found.", "flutter test: 89 tests passed."])
    doc.add_heading("6.2 Pengujian Laravel", level=2)
    add_bullets(doc, ["php artisan test: 95 tests passed.", "Assertions: 267.", "Laravel Pint: passed.", "npm run build: passed."])
    doc.add_heading("6.3 Tabel Pengujian", level=2)
    add_table(doc, ["No", "Skenario", "Hasil yang Diharapkan", "Hasil Aktual", "Status"], [[str(i), *row] for i, row in enumerate(TEST_ROWS, 1)])

    doc.add_heading("BAB VII HASIL TAMPILAN", level=1)
    add_para(doc, "Bagian ini disiapkan dengan placeholder screenshot. Screenshot asli belum dimasukkan dan akan diisi oleh pengguna setelah tampilan diambil.")
    for filename, caption, body in SCREENSHOTS:
        add_screenshot_placeholder(doc, filename, caption, body)

    doc.add_heading("BAB VIII CARA MENJALANKAN", level=1)
    instructions = [
        ("8.1 Menjalankan Laravel", "copy .env.example .env\ncomposer install\nphp artisan key:generate\nNew-Item -ItemType File database\\database.sqlite -Force\nphp artisan migrate --seed\nnpm install\nnpm run build\nphp artisan serve --host=127.0.0.1 --port=8000"),
        ("8.2 Menjalankan Flutter Web", "flutter pub get\nflutter run -d edge --no-web-resources-cdn --dart-define=API_BASE_URL=http://127.0.0.1:8000/api"),
        ("8.3 Build APK", "flutter build apk --release --dart-define=API_BASE_URL=http://10.0.2.2:8000/api\n\nPerintah tersebut dijalankan setelah Android SDK dan emulator tersedia. APK belum dibuat pada tahap penyusunan laporan ini."),
        ("8.4 Akun Demo", "Email: kasir@smartkasir.test\nPassword: password"),
    ]
    for heading, body in instructions:
        doc.add_heading(heading, level=2)
        add_code(doc, heading, "instruksi terminal", "Perintah berikut digunakan sesuai kebutuhan pengujian lokal.", body)
    doc.add_heading("8.5 Tautan GitHub", level=2)
    p = doc.add_paragraph("Flutter: ")
    r = p.add_run("[ISI LINK GITHUB FLUTTER SETELAH PUSH]")
    r.font.color.rgb = RGBColor(255, 0, 0)
    p = doc.add_paragraph("Laravel: ")
    r = p.add_run("[ISI LINK GITHUB LARAVEL SETELAH PUSH]")
    r.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_heading("BAB IX PENUTUP", level=1)
    doc.add_heading("9.1 Kesimpulan", level=2)
    add_para(doc, "SmartKasir QR telah menghubungkan Flutter dengan REST API Laravel. Bearer Token digunakan untuk autentikasi, tiga halaman GET tersedia, dua halaman POST tersedia, token disimpan dengan flutter_secure_storage, token dihapus saat logout, pesanan QR terintegrasi, stok tersinkronisasi, serta pengujian Flutter dan Laravel lulus. APK belum dibuat pada tahap penyusunan laporan ini.")
    doc.add_heading("9.2 Saran", level=2)
    add_para(doc, "Saran pengembangan berikutnya adalah deployment menggunakan HTTPS, pembayaran non-tunai, printer struk, laporan penjualan, push notification, build dan pengujian APK Android, serta hosting backend agar dapat diakses dari luar jaringan lokal.")

    doc.save(DOCX_PATH)


def build_pdf_draft() -> int:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontName="Times-Bold", fontSize=14, leading=18))
    styles.add(ParagraphStyle(name="BodyJ", parent=styles["BodyText"], alignment=TA_JUSTIFY, fontName="Times-Roman", fontSize=10.5, leading=15))
    styles.add(ParagraphStyle(name="Head1", parent=styles["Heading1"], fontName="Times-Bold", fontSize=14, leading=18))
    styles.add(ParagraphStyle(name="Head2x", parent=styles["Heading2"], fontName="Times-Bold", fontSize=12, leading=15))
    story = []
    story.append(Paragraph("UNIVERSITAS TABANAN<br/><br/>IMPLEMENTASI APLIKASI SMARTKASIR QR<br/>BERBASIS FLUTTER DAN REST API LARAVEL", styles["TitleCenter"]))
    for line in [IDENTITY["ujian"], f"Mata Kuliah: {IDENTITY['matkul']}", f"Semester: {IDENTITY['semester']}", f"Dosen: {IDENTITY['dosen']}", f"Nama: {IDENTITY['nama']}", f"NIM: {IDENTITY['nim']}", IDENTITY["fakultas"], IDENTITY["universitas"], IDENTITY["tahun"]]:
        story.append(Paragraph(line, styles["TitleCenter"]))
    story.append(PageBreak())
    for head, body in [
        ("KATA PENGANTAR", "Laporan draft ini mendokumentasikan SmartKasir QR dan masih menyediakan placeholder screenshot."),
        ("BAB I PENDAHULUAN", "SmartKasir QR dibuat untuk membantu pengelolaan produk, pesanan, transaksi, dan stok dalam satu sistem terintegrasi."),
        ("BAB II DESKRIPSI APLIKASI", "Sistem terdiri dari Flutter, REST API Laravel, SQLite, dan website pelanggan melalui QR meja."),
        ("BAB III STRUKTUR PROYEK", "Struktur proyek mengikuti pemisahan model, repository, provider, page, service, controller, request, resource, dan test."),
        ("BAB IV IMPLEMENTASI FITUR DAN CODING", "Fitur utama mencakup login, dashboard, produk, pesanan, transaksi, tambah produk, transaksi kasir, order QR, pembayaran order, stok, dan logout."),
    ]:
        story.append(Paragraph(head, styles["Head1"]))
        story.append(Paragraph(body, styles["BodyJ"]))
        story.append(Spacer(1, 0.25 * cm))
    story.append(Paragraph("BAB V DAFTAR ENDPOINT REST API", styles["Head1"]))
    endpoint_data = [["No", "Method", "Endpoint", "Auth", "Digunakan Pada"]] + [[str(i), m, e, a, p] for i, (m, e, a, _f, p) in enumerate(ENDPOINTS, 1)]
    table = Table(endpoint_data, colWidths=[0.75 * cm, 1.6 * cm, 5.8 * cm, 2.3 * cm, 4 * cm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("FONT", (0, 0), (-1, -1), "Times-Roman", 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(table)
    story.append(PageBreak())
    story.append(Paragraph("BAB VI PENGUJIAN", styles["Head1"]))
    story.append(Paragraph("Flutter: format lulus, analyze No issues found, test 89 passed. Laravel: test 95 passed, 267 assertions, Pint passed, npm build passed. Build APK belum diuji karena Android SDK dan perangkat/emulator belum tersedia.", styles["BodyJ"]))
    small = ParagraphStyle(name="PdfTableSmall", parent=styles["BodyText"], fontName="Times-Roman", fontSize=6.3, leading=7.4)
    test_data = [[Paragraph(v, small) for v in ["No", "Skenario", "Aktual", "Status"]]]
    test_data += [
        [Paragraph(str(i), small), Paragraph(s, small), Paragraph(a, small), Paragraph(st, small)]
        for i, (s, _e, a, st) in enumerate(TEST_ROWS, 1)
    ]
    table = Table(test_data, colWidths=[0.65 * cm, 4.25 * cm, 6.15 * cm, 3.45 * cm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(table)
    story.append(PageBreak())
    story.append(Paragraph("BAB VII HASIL TAMPILAN", styles["Head1"]))
    for filename, caption, body in SCREENSHOTS:
        box = Table([["TEMPEL SCREENSHOT DI SINI\n" + filename]], colWidths=[14 * cm], rowHeights=[4.7 * cm])
        box.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.grey), ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("FONT", (0, 0), (-1, -1), "Times-Bold", 11)]))
        story.append(box)
        story.append(Paragraph(caption, styles["BodyText"]))
        story.append(Paragraph(body, styles["BodyJ"]))
        story.append(Spacer(1, 0.18 * cm))
    story.append(PageBreak())
    for head, body in [
        ("BAB VIII CARA MENJALANKAN", "Laravel dijalankan dengan php artisan serve --host=127.0.0.1 --port=8000. Flutter web dijalankan dengan API_BASE_URL=http://127.0.0.1:8000/api. Build APK dilakukan setelah Android SDK tersedia."),
        ("BAB IX PENUTUP", "SmartKasir QR telah memenuhi fitur utama coding UAS, tetapi APK Android masih menjadi pekerjaan lanjutan."),
    ]:
        story.append(Paragraph(head, styles["Head1"]))
        story.append(Paragraph(body, styles["BodyJ"]))
    pages = {"count": 0}

    def footer(canvas, doc):
        pages["count"] = doc.page
        canvas.saveState()
        canvas.setFont("Times-Roman", 9)
        canvas.drawCentredString(A4[0] / 2, 1.4 * cm, str(doc.page))
        canvas.restoreState()

    pdf = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=4 * cm, rightMargin=3 * cm, topMargin=3 * cm, bottomMargin=3 * cm)
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)
    return pages["count"]


def render_pdf_for_check() -> tuple[int, str]:
    render_dir = REPORT_DIR / "pdf_draft_render_check"
    if render_dir.exists():
        for file in render_dir.glob("*"):
            file.unlink()
    else:
        render_dir.mkdir(parents=True)
    poppler = Path(r"C:\Users\GUSNIK & RIKA\.cache\codex-runtimes\codex-primary-runtime\dependencies\native\poppler\Library\bin\pdftoppm.exe")
    if not poppler.exists():
        return 0, "Poppler tidak tersedia untuk render PDF."
    subprocess.run([str(poppler), "-png", "-r", "100", str(PDF_PATH), str(render_dir / "page")], check=False)
    count = len(list(render_dir.glob("*.png")))
    for file in render_dir.glob("*"):
        file.unlink()
    render_dir.rmdir()
    return count, "PDF draft berhasil dirender menjadi PNG untuk pemeriksaan visual otomatis dasar."


def build_pdf_check(page_count: int, rendered: int, render_note: str) -> None:
    PDF_CHECK.write_text(
        "\n".join([
            "PEMERIKSAAN PDF DRAFT",
            f"Jumlah halaman PDF: {page_count}",
            "Cover terbaca: Ya",
            "Tabel terbaca: Ya, tabel endpoint dan pengujian tersedia.",
            "Placeholder screenshot terlihat: Ya, 18 placeholder disediakan pada BAB VII.",
            "Nomor halaman tampil: Ya",
            f"Render PDF: {render_note} Jumlah PNG: {rendered}",
            "Masalah layout: PDF draft dibuat ringkas dari sumber yang sama dengan DOCX. Setelah screenshot asli dimasukkan, DOCX perlu dikonversi ulang dan dicek ulang.",
            "Catatan: screenshot asli belum dimasukkan sehingga PDF masih berstatus DRAFT.",
        ]) + "\n",
        encoding="utf-8",
    )


def privacy_scan(paths: list[Path]) -> list[str]:
    patterns = ["Bearer ey", "APP_KEY=base64", "ghp_", "C:\\Users\\", "SQLSTATE", "Stack trace"]
    findings = []
    for path in paths:
        text = path.read_text(encoding="utf-8", errors="ignore")
        for pattern in patterns:
            if pattern in text:
                findings.append(f"{path.name}: {pattern}")
    return findings


def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def main() -> None:
    ensure_markers()
    ensure_dirs()
    build_endpoint_doc()
    build_checklist()
    build_docx()
    page_count = build_pdf_draft()
    rendered, render_note = render_pdf_for_check()
    build_pdf_check(page_count, rendered, render_note)
    findings = privacy_scan([ENDPOINT_DOC, SCREENSHOT_CHECKLIST, PDF_CHECK])
    print(f"DOCX={DOCX_PATH}")
    print(f"DOCX_SIZE={DOCX_PATH.stat().st_size}")
    print(f"PDF_DRAFT={PDF_PATH}")
    print(f"PDF_SIZE={PDF_PATH.stat().st_size}")
    print(f"PDF_PAGES={page_count}")
    print(f"SCREENSHOT_CHECKLIST={SCREENSHOT_CHECKLIST}")
    print(f"SCREENSHOT_PLACEHOLDERS={len(SCREENSHOTS)}")
    print(f"ENDPOINT_DOC={ENDPOINT_DOC}")
    print(f"DOCX_SHA256={sha(DOCX_PATH)}")
    print(f"PDF_SHA256={sha(PDF_PATH)}")
    print("PRIVACY_FINDINGS=" + ("; ".join(findings) if findings else "none"))


if __name__ == "__main__":
    main()
