from __future__ import annotations

import hashlib
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"C:\SmartKasirQR")
API = ROOT / "smartkasir-qr-api"
MOBILE = ROOT / "smartkasir-qr-mobile"
SUBMISSION = ROOT / "submission"
LAPORAN = SUBMISSION / "03_Laporan"
API_DOC = SUBMISSION / "05_Dokumentasi_API"
APK_DIR = SUBMISSION / "02_APK"
SHOT_DIR = SUBMISSION / "04_Screenshot"

NOW = datetime.now()
FINAL_DATE = NOW.strftime("%d %B %Y")

ENDPOINTS = [
    (1, "POST", "/api/login", "Tidak", "Login kasir dan membuat Bearer Token.", "Login Flutter"),
    (2, "GET", "/api/me", "Bearer Token", "Mengambil profil sesi aktif.", "Pemulihan sesi"),
    (3, "POST", "/api/logout", "Bearer Token", "Logout dan hapus token server.", "Profil"),
    (4, "GET", "/api/dashboard", "Bearer Token", "Ringkasan dashboard kasir.", "Beranda"),
    (5, "GET", "/api/categories", "Bearer Token", "Daftar kategori produk.", "Produk/Tambah Produk"),
    (6, "GET", "/api/products", "Bearer Token", "Daftar produk dengan filter.", "Produk"),
    (7, "GET", "/api/products/{id}", "Bearer Token", "Detail produk.", "Detail Produk"),
    (8, "POST", "/api/products", "Bearer Token", "Tambah produk baru.", "Tambah Produk"),
    (9, "GET", "/api/orders", "Bearer Token", "Daftar pesanan QR.", "Pesanan"),
    (10, "GET", "/api/orders/{id}", "Bearer Token", "Detail pesanan.", "Detail Pesanan"),
    (11, "PATCH", "/api/orders/{id}/confirm", "Bearer Token", "Konfirmasi dan kurangi stok.", "Detail Pesanan"),
    (12, "PATCH", "/api/orders/{id}/process", "Bearer Token", "Ubah pesanan menjadi diproses.", "Detail Pesanan"),
    (13, "PATCH", "/api/orders/{id}/ready", "Bearer Token", "Tandai pesanan siap.", "Detail Pesanan"),
    (14, "PATCH", "/api/orders/{id}/cancel", "Bearer Token", "Batalkan pesanan dan kembalikan stok bila perlu.", "Detail Pesanan"),
    (15, "GET", "/api/transactions", "Bearer Token", "Daftar transaksi.", "Transaksi"),
    (16, "GET", "/api/transactions/{id}", "Bearer Token", "Detail transaksi.", "Detail Transaksi"),
    (17, "POST", "/api/transactions", "Bearer Token", "Transaksi kasir atau pembayaran order QR.", "Transaksi/Pembayaran"),
    (18, "GET", "/api/public/tables/{qrToken}/menu", "Tidak", "Menu publik pelanggan QR.", "Website QR"),
    (19, "POST", "/api/public/tables/{qrToken}/orders", "Tidak", "Kirim pesanan pelanggan QR.", "Website QR"),
    (20, "GET", "/api/public/orders/{orderNumber}/status", "Tidak", "Cek status pesanan publik.", "Website QR"),
]

TEST_ROWS = [
    ("Login benar", "Kasir berhasil masuk dengan akun demo.", "Lulus pada test Flutter dan API Laravel.", "Lulus"),
    ("Login salah", "Validasi error tampil.", "Lulus pada test auth.", "Lulus"),
    ("Pemulihan sesi", "Token dicek melalui GET /api/me.", "Lulus pada test Flutter.", "Lulus"),
    ("Logout", "Token lokal dihapus dan halaman kembali login.", "Lulus pada test Flutter.", "Lulus"),
    ("GET Dashboard", "Dashboard tampil dari API.", "Lulus pada test dashboard.", "Lulus"),
    ("GET Produk", "Daftar produk tampil.", "Lulus pada test produk.", "Lulus"),
    ("GET Pesanan", "Daftar pesanan tampil.", "Lulus pada test pesanan.", "Lulus"),
    ("GET Transaksi", "Daftar transaksi tampil.", "Lulus pada test transaksi.", "Lulus"),
    ("POST Produk", "Produk baru tersimpan.", "Lulus pada test ProductApiTest.", "Lulus"),
    ("SKU duplikat", "Ditolak validasi server.", "Lulus pada test ProductApiTest.", "Lulus"),
    ("POST Transaksi", "Transaksi tersimpan dan stok berkurang.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Pembayaran kurang", "Ditolak validasi server.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Pesanan QR pending", "Order publik dibuat pending.", "Lulus pada API/web test.", "Lulus"),
    ("Konfirmasi pesanan", "Status confirmed dan stok berkurang.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Proses pesanan", "Status processing.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Pesanan siap", "Status ready.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Pembayaran order", "Order completed dan transaksi dibuat.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Pembayaran ganda", "Ditolak 422.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Cancel pending", "Order cancelled tanpa stok berubah.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("Cancel confirmed", "Order cancelled dan stok kembali.", "Lulus pada StageThreeApiTest.", "Lulus"),
    ("401 Unauthorized", "Provider menghapus sesi.", "Lulus pada test Flutter.", "Lulus"),
    ("Server tidak tersedia", "UI menampilkan error.", "Lulus pada widget/provider test.", "Lulus"),
    ("Build APK release", "APK release dibuat.", "Belum diuji karena Android SDK tidak ditemukan.", "Belum Lulus"),
    ("Instalasi APK", "APK terpasang di device.", "Belum diuji karena tidak ada Android device/emulator.", "Belum Lulus"),
    ("Login dari APK", "Login berhasil dari APK.", "Belum diuji.", "Belum Lulus"),
    ("GET dari APK", "GET API berhasil dari APK.", "Belum diuji.", "Belum Lulus"),
    ("POST dari APK", "POST API berhasil dari APK.", "Belum diuji.", "Belum Lulus"),
    ("Logout dari APK", "Logout berhasil dari APK.", "Belum diuji.", "Belum Lulus"),
]


def ensure_dirs() -> None:
    for path in [SUBMISSION, LAPORAN, API_DOC, APK_DIR, SHOT_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def read_snippet(path: Path, needles: list[str], max_lines: int = 34) -> str:
    text = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    start = 0
    for index, line in enumerate(text):
        if any(needle in line for needle in needles):
            start = max(0, index - 4)
            break
    return "\n".join(text[start : start + max_lines])


def write_text_files() -> None:
    endpoint_lines = [
        "# Dokumentasi Endpoint API SmartKasir QR",
        "",
        "Base URL lokal browser: `http://127.0.0.1:8000/api`",
        "",
        "| No | Method | Endpoint | Autentikasi | Fungsi | Halaman Pengguna |",
        "|---:|---|---|---|---|---|",
    ]
    for row in ENDPOINTS:
        endpoint_lines.append(f"| {row[0]} | {row[1]} | `{row[2]}` | {row[3]} | {row[4]} | {row[5]} |")
    (API_DOC / "ENDPOINT_API.md").write_text("\n".join(endpoint_lines) + "\n", encoding="utf-8")

    (SUBMISSION / "README_KONEKSI_SERVER.txt").write_text(
        "\n".join(
            [
                "SMARTKASIR QR - README KONEKSI SERVER",
                "",
                "Jalankan Laravel dari folder C:\\SmartKasirQR\\smartkasir-qr-api:",
                "php artisan serve --host=0.0.0.0 --port=8000",
                "",
                "URL API untuk Android Emulator:",
                "http://10.0.2.2:8000/api",
                "",
                "URL API untuk perangkat fisik satu Wi-Fi:",
                "http://IPV4-KOMPUTER:8000/api",
                "Cari IPv4 komputer dengan perintah ipconfig.",
                "",
                "Alternatif perangkat fisik via USB:",
                "adb reverse tcp:8000 tcp:8000",
                "Gunakan API URL: http://127.0.0.1:8000/api",
                "",
                "Akun demo:",
                "Email: kasir@smartkasir.test",
                "Password: password",
                "",
                "Catatan: Windows Firewall harus mengizinkan php.exe menerima koneksi lokal.",
                "HTTP cleartext hanya untuk demo lokal. Deployment publik sebaiknya memakai HTTPS.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    (SUBMISSION / "README_PENGUMPULAN.txt").write_text(
        "\n".join(
            [
                "SMARTKASIR QR - README PENGUMPULAN",
                "",
                "Nama proyek: SmartKasir QR",
                "Nama mahasiswa: [ISI NAMA LENGKAP MAHASISWA]",
                "NIM: 2323050026",
                "Program Studi: Sistem Informasi",
                "Universitas: Universitas Tabanan",
                "",
                "Deskripsi:",
                "SmartKasir QR adalah aplikasi kasir Flutter yang terhubung dengan REST API Laravel dan website QR meja.",
                "",
                "Struktur:",
                "01_Source: source Flutter dan Laravel bersih",
                "02_APK: APK release jika Android SDK/device sudah tersedia",
                "03_Laporan: DOCX/PDF laporan UAS",
                "04_Screenshot: screenshot aktual",
                "05_Dokumentasi_API: dokumentasi endpoint",
                "",
                "Akun demo:",
                "Email: kasir@smartkasir.test",
                "Password: password",
                "",
                "Menjalankan Laravel:",
                "cd C:\\SmartKasirQR\\smartkasir-qr-api",
                "composer install",
                "copy .env.example .env",
                "php artisan key:generate",
                "New-Item -ItemType File database\\database.sqlite -Force",
                "php artisan migrate --seed",
                "php artisan serve --host=0.0.0.0 --port=8000",
                "",
                "Menjalankan Flutter Web:",
                "cd C:\\SmartKasirQR\\smartkasir-qr-mobile",
                "flutter pub get",
                "flutter run -d edge --dart-define=API_BASE_URL=http://127.0.0.1:8000/api",
                "",
                "Build APK:",
                "flutter build apk --release --dart-define=API_BASE_URL=http://10.0.2.2:8000/api",
                "",
                "Status test terbaru:",
                "Flutter analyze: No issues found.",
                "Flutter test: 89 tests passed.",
                "Laravel test: 95 tests passed, 267 assertions.",
                "Laravel Pint: passed.",
                "npm run build: passed.",
                "",
                "Tautan GitHub Flutter: [ISI LINK GITHUB FLUTTER SETELAH PUSH]",
                "Tautan GitHub Laravel: [ISI LINK GITHUB LARAVEL SETELAH PUSH]",
                "",
                "Catatan Tahap 11:",
                "Android SDK/device belum tersedia pada audit ini, sehingga APK release dan uji APK belum selesai.",
                f"Tanggal finalisasi dokumen: {FINAL_DATE}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    (LAPORAN / "IDENTITAS_BELUM_LENGKAP.txt").write_text(
        "Field identitas yang belum ditemukan dari source: Nama lengkap mahasiswa.\n"
        "Gunakan placeholder [ISI NAMA LENGKAP MAHASISWA] pada laporan.\n"
        "Placeholder GitHub juga disediakan karena push dilakukan pengguna.\n",
        encoding="utf-8",
    )
    (APK_DIR / "APK_BELUM_DIBUAT_KARENA_ANDROID_SDK.txt").write_text(
        "APK release belum dibuat karena flutter doctor tidak menemukan Android SDK dan tidak ada Android device/emulator.\n"
        "Setelah Android SDK tersedia, jalankan build sesuai README_KONEKSI_SERVER.txt.\n",
        encoding="utf-8",
    )
    (APK_DIR / "BUILD_INFO.txt").write_text(
        "\n".join(
            [
                "Nama aplikasi: SmartKasir QR",
                "Application ID: com.gusnik.smartkasirqr",
                "Version name: 1.0.0",
                "Version code: 1",
                "Flutter version: 3.35.6",
                "Dart version: 3.9.2",
                f"Tanggal audit: {FINAL_DATE}",
                "Build mode: release belum dibuat",
                "API base URL rencana emulator: http://10.0.2.2:8000/api",
                "Device pengujian: belum tersedia",
                "SHA-256 APK: belum tersedia",
                "Catatan: backend Laravel harus aktif saat aplikasi digunakan.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    (SHOT_DIR / "SCREENSHOT_BELUM_LENGKAP.txt").write_text(
        "Screenshot APK Android release belum tersedia karena Android SDK/device belum tersedia.\n"
        "Screenshot website QR dapat diambil dari browser setelah server Laravel aktif.\n",
        encoding="utf-8",
    )


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold


def add_code(doc: Document, title: str, code: str) -> None:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("IMPLEMENTASI APLIKASI SMARTKASIR QR\nBERBASIS FLUTTER DAN REST API LARAVEL")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph("")
    for line in [
        "Ujian Akhir Semester Genap Tahun Akademik 2025/2026",
        "Mata Kuliah: Pemrograman Mobile Lanjutan",
        "Semester: VI",
        "Dosen: Luh Ayu Diah Fernita Sari, S.Kom., M.T.",
        "NIM: 2323050026",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Nama: ")
    missing = p.add_run("[ISI NAMA LENGKAP MAHASISWA]")
    missing.font.color.rgb = RGBColor(255, 0, 0)
    missing.bold = True
    for line in ["PROGRAM STUDI SISTEM INFORMASI", "FAKULTAS SAINS DAN TEKNOLOGI", "UNIVERSITAS TABANAN", "2026"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("KATA PENGANTAR", level=1)
    add_docx_paragraph(doc, "Puji syukur penulis panjatkan karena laporan UAS SmartKasir QR ini dapat disusun sebagai dokumentasi implementasi aplikasi kasir berbasis Flutter dan REST API Laravel. Laporan ini menjelaskan fitur, endpoint, pengujian, serta catatan rilis yang ditemukan selama audit akhir.")
    doc.add_page_break()

    doc.add_heading("DAFTAR ISI", level=1)
    for item in [
        "BAB I PENDAHULUAN",
        "BAB II DESKRIPSI APLIKASI",
        "BAB III STRUKTUR PROYEK",
        "BAB IV IMPLEMENTASI FITUR DAN CODING",
        "BAB V ENDPOINT REST API",
        "BAB VI PENGUJIAN",
        "BAB VII HASIL TAMPILAN",
        "BAB VIII CARA MENJALANKAN",
        "BAB IX PENUTUP",
        "LAMPIRAN",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    chapters = [
        ("BAB I PENDAHULUAN", [
            ("1.1 Latar Belakang", "Kasir usaha kecil membutuhkan pencatatan transaksi yang rapi, pengelolaan stok, dan kanal pemesanan pelanggan yang cepat. SmartKasir QR dibuat untuk menghubungkan aplikasi kasir Flutter dengan backend Laravel serta website pemesanan QR meja."),
            ("1.2 Rumusan Masalah", "Bagaimana membangun aplikasi kasir yang memiliki login, GET API, POST API, pembayaran, sinkronisasi stok, dan website QR pelanggan."),
            ("1.3 Tujuan", "Membuat aplikasi SmartKasir QR yang dapat dipakai kasir untuk mengelola produk, pesanan, dan transaksi melalui REST API Laravel."),
            ("1.4 Manfaat", "Aplikasi membantu kasir melihat pesanan QR, mengurangi kesalahan pencatatan, dan menjaga stok tetap sinkron."),
            ("1.5 Batasan Sistem", "Sistem memakai SQLite untuk UAS dan HTTP lokal untuk demonstrasi. Deployment publik sebaiknya memakai HTTPS."),
        ]),
        ("BAB II DESKRIPSI APLIKASI", [
            ("2.1 Gambaran SmartKasir QR", "SmartKasir QR terdiri dari aplikasi Flutter kasir, REST API Laravel, database SQLite, dan website pelanggan yang diakses melalui QR meja."),
            ("2.2 Pengguna Sistem", "Pengguna utama adalah kasir. Pelanggan memakai website QR publik tanpa login."),
            ("2.3 Teknologi", "Flutter, Dart, Dio, Provider, flutter_secure_storage, Laravel, Sanctum, SQLite, Blade, Vite, dan PHPUnit."),
            ("2.4 Arsitektur Sistem", "Flutter berkomunikasi dengan Laravel memakai Bearer Token. Laravel menjadi sumber kebenaran untuk stok, total transaksi, kembalian, dan status pesanan."),
            ("2.5 Alur Pesanan QR", "Pelanggan membuat order pending. Kasir mengonfirmasi, memproses, menandai siap, lalu menerima pembayaran hingga order completed."),
        ]),
        ("BAB III STRUKTUR PROYEK", [
            ("3.1 Struktur Flutter", "Folder lib berisi config, core, navigation, dan features untuk auth, dashboard, products, orders, transactions, dan profile."),
            ("3.2 Struktur Laravel", "Folder app berisi controller API/web, request, resource, model, service, dan support. Route API ada pada routes/api.php dan website QR ada pada routes/web.php."),
            ("3.3 Struktur Database", "Database SQLite memakai tabel user, personal access token, categories, products, restaurant tables, orders, order items, transactions, dan transaction items."),
        ]),
        ("BAB IV IMPLEMENTASI FITUR DAN CODING", [
            ("4.1 Login dan Bearer Token", "Login dilakukan melalui POST /api/login. Token disimpan aman dan dipakai otomatis oleh ApiClient."),
            ("4.2 Penyimpanan Token", "Token memakai flutter_secure_storage melalui TokenStorage."),
            ("4.3 Home Dashboard", "Dashboard memakai GET /api/dashboard."),
            ("4.4 Bottom Navigation", "Navigasi utama memiliki Beranda, Pesanan, Produk, dan Transaksi."),
            ("4.5 GET Produk", "Produk mendukung daftar, detail, pencarian, filter kategori, dan filter ketersediaan."),
            ("4.6 GET Pesanan", "Pesanan mendukung daftar, detail, filter status, dan aksi status."),
            ("4.7 GET Transaksi", "Transaksi mendukung daftar, detail, pencarian, dan filter tanggal."),
            ("4.8 POST Tambah Produk", "Flutter hanya mengirim field produk yang diperlukan; validasi final ada di Laravel."),
            ("4.9 POST Transaksi Kasir", "Flutter mengirim item dan paid amount. Backend menghitung harga, total, kembalian, dan stok."),
            ("4.10 Pemesanan Pelanggan QR", "Website QR mengirim product_id, quantity, dan notes. Backend menghitung total dan membuat order pending."),
            ("4.11 Perubahan Status Pesanan", "PATCH confirm/process/ready/cancel mengubah status sesuai aturan."),
            ("4.12 Pembayaran Pesanan", "Pembayaran order QR memakai POST /api/transactions dengan order_id."),
            ("4.13 Sinkronisasi Stok", "Stok berkurang saat confirm order atau transaksi langsung, dan tidak berkurang dua kali saat pembayaran order."),
            ("4.14 Logout dan Reset State", "Logout menghapus token serta reset provider."),
            ("4.15 Loading dan Error State", "Setiap halaman utama memiliki loading, empty, error, retry, search/filter, dan refresh."),
        ]),
    ]
    for chapter, sections in chapters:
        doc.add_heading(chapter, level=1)
        for heading, body in sections:
            doc.add_heading(heading, level=2)
            add_docx_paragraph(doc, body)

    snippets = [
        ("ApiClient Bearer Token", MOBILE / "lib/core/network/api_client.dart", ["Authorization"]),
        ("AuthRepository Login", MOBILE / "lib/features/auth/repositories/auth_repository.dart", ["Future<UserModel> login", "login("]),
        ("TokenStorage", MOBILE / "lib/core/storage/token_storage.dart", ["class TokenStorage"]),
        ("AuthProvider", MOBILE / "lib/features/auth/providers/auth_provider.dart", ["class AuthProvider"]),
        ("DashboardRepository", MOBILE / "lib/features/dashboard/repositories/dashboard_repository.dart", ["fetchDashboard"]),
        ("ProductRepository GET/POST", MOBILE / "lib/features/products/repositories/product_repository.dart", ["fetchProducts", "createProduct"]),
        ("OrderRepository PATCH", MOBILE / "lib/features/orders/repositories/order_repository.dart", ["confirmOrder"]),
        ("TransactionRepository POST", MOBILE / "lib/features/transactions/repositories/transaction_repository.dart", ["createTransaction", "createOrderPayment"]),
        ("MainNavigationPage", MOBILE / "lib/navigation/main_navigation_page.dart", ["class MainNavigationPage"]),
        ("routes/api.php", API / "routes/api.php", ["Route::post('/login'"]),
        ("OrderService Stok", API / "app/Services/OrderService.php", ["public function confirm"]),
        ("TransactionService", API / "app/Services/TransactionService.php", ["public function create"]),
    ]
    for title, path, needles in snippets:
        add_code(doc, title, read_snippet(path, needles))

    doc.add_heading("BAB V ENDPOINT REST API", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for idx, head in enumerate(["No", "Method", "Endpoint", "Auth", "Fungsi", "Halaman"]):
        table.rows[0].cells[idx].text = head
    for row in ENDPOINTS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)

    doc.add_heading("BAB VI PENGUJIAN", level=1)
    add_docx_paragraph(doc, "Validasi terakhir: flutter analyze No issues found, flutter test 89 tests passed, php artisan test 95 tests passed dengan 267 assertions, Laravel Pint passed, dan npm run build passed.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, head in enumerate(["No", "Skenario", "Diharapkan", "Aktual", "Status"]):
        table.rows[0].cells[idx].text = head
    for i, row in enumerate(TEST_ROWS, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        for idx, value in enumerate(row, start=1):
            cells[idx].text = value

    doc.add_heading("BAB VII HASIL TAMPILAN", level=1)
    add_docx_paragraph(doc, "Screenshot APK Android release belum tersedia karena Android SDK dan device/emulator belum tersedia pada audit ini. Screenshot harus diambil ulang dari APK release setelah toolchain Android siap.")

    doc.add_heading("BAB VIII CARA MENJALANKAN", level=1)
    for heading, body in [
        ("8.1 Laravel dan SQLite", "Jalankan composer install, salin .env.example, buat database/database.sqlite, php artisan migrate --seed, lalu php artisan serve --host=0.0.0.0 --port=8000."),
        ("8.2 Flutter Web", "Jalankan flutter run -d edge --dart-define=API_BASE_URL=http://127.0.0.1:8000/api."),
        ("8.3 Build APK", "Setelah Android SDK tersedia, jalankan flutter build apk --release --dart-define=API_BASE_URL=http://10.0.2.2:8000/api untuk emulator."),
        ("8.4 Instal APK", "Gunakan adb install -r SmartKasirQR-v1.0.0.apk setelah APK tersedia."),
        ("8.5 Akun Demo", "Email kasir@smartkasir.test dan password password."),
        ("8.6 Tautan GitHub", "Flutter: [ISI LINK GITHUB FLUTTER SETELAH PUSH]. Laravel: [ISI LINK GITHUB LARAVEL SETELAH PUSH]."),
    ]:
        doc.add_heading(heading, level=2)
        add_docx_paragraph(doc, body)

    doc.add_heading("BAB IX PENUTUP", level=1)
    doc.add_heading("9.1 Kesimpulan", level=2)
    add_docx_paragraph(doc, "Fitur coding utama SmartKasir QR telah terpenuhi dan validasi source lulus. Tahap rilis Android belum lengkap karena Android SDK/device belum tersedia.")
    doc.add_heading("9.2 Saran", level=2)
    add_docx_paragraph(doc, "Pengembangan berikutnya adalah menyiapkan Android SDK, melakukan uji APK release di perangkat nyata, mengambil screenshot APK, dan mengganti placeholder identitas serta GitHub.")

    path = LAPORAN / "Laporan_UAS_SmartKasir_QR.docx"
    doc.save(path)
    return path


def build_pdf() -> Path:
    pdf = LAPORAN / "Laporan_UAS_SmartKasir_QR.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, fontName="Times-Bold", fontSize=14, leading=18))
    styles.add(ParagraphStyle(name="BodyJustify", parent=styles["BodyText"], alignment=TA_JUSTIFY, fontName="Times-Roman", fontSize=11, leading=16))
    styles.add(ParagraphStyle(name="Head", parent=styles["Heading1"], fontName="Times-Bold", fontSize=14, leading=18, spaceAfter=8))
    styles.add(ParagraphStyle(name="SubHead", parent=styles["Heading2"], fontName="Times-Bold", fontSize=12, leading=15, spaceAfter=6))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontName="Times-Roman", fontSize=8, leading=10))
    story = []
    story.append(Paragraph("IMPLEMENTASI APLIKASI SMARTKASIR QR<br/>BERBASIS FLUTTER DAN REST API LARAVEL", styles["CenterTitle"]))
    story += [Spacer(1, 1 * cm)]
    for line in [
        "Ujian Akhir Semester Genap Tahun Akademik 2025/2026",
        "Mata Kuliah: Pemrograman Mobile Lanjutan",
        "Semester: VI",
        "Dosen: Luh Ayu Diah Fernita Sari, S.Kom., M.T.",
        "NIM: 2323050026",
        'Nama: <font color="red">[ISI NAMA LENGKAP MAHASISWA]</font>',
        "PROGRAM STUDI SISTEM INFORMASI",
        "FAKULTAS SAINS DAN TEKNOLOGI",
        "UNIVERSITAS TABANAN",
        "2026",
    ]:
        story.append(Paragraph(line, styles["CenterTitle"]))
    story.append(PageBreak())

    sections = [
        ("KATA PENGANTAR", "Laporan ini disusun untuk mendokumentasikan implementasi SmartKasir QR berbasis Flutter dan REST API Laravel."),
        ("BAB I PENDAHULUAN", "SmartKasir QR menjawab kebutuhan kasir untuk mengelola produk, pesanan QR, transaksi, dan stok secara terintegrasi."),
        ("BAB II DESKRIPSI APLIKASI", "Aplikasi terdiri dari Flutter kasir, REST API Laravel, SQLite, dan website pelanggan melalui QR meja."),
        ("BAB III STRUKTUR PROYEK", "Source Flutter berada pada smartkasir-qr-mobile dan backend Laravel pada smartkasir-qr-api."),
        ("BAB IV IMPLEMENTASI FITUR DAN CODING", "Fitur utama meliputi login Bearer Token, dashboard, produk, pesanan, transaksi, tambah produk, transaksi kasir, pembayaran order QR, dan logout."),
    ]
    for head, body in sections:
        story.append(Paragraph(head, styles["Head"]))
        story.append(Paragraph(body, styles["BodyJustify"]))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("BAB V ENDPOINT REST API", styles["Head"]))
    data = [["No", "Method", "Endpoint", "Auth", "Halaman"]]
    data += [[str(n), method, endpoint, auth, page] for n, method, endpoint, auth, _func, page in ENDPOINTS]
    table = Table(data, colWidths=[0.8 * cm, 1.8 * cm, 6 * cm, 2.4 * cm, 4 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONT", (0, 0), (-1, -1), "Times-Roman", 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.append(PageBreak())

    story.append(Paragraph("BAB VI PENGUJIAN", styles["Head"]))
    story.append(Paragraph("Flutter analyze lulus, Flutter test 89 test lulus, Laravel test 95 test dan 267 assertion lulus, Laravel Pint lulus, dan npm run build lulus. APK belum diuji karena Android SDK/device belum tersedia.", styles["BodyJustify"]))
    data = [["No", "Skenario", "Aktual", "Status"]]
    data += [[str(i), s, actual, status] for i, (s, _exp, actual, status) in enumerate(TEST_ROWS, start=1)]
    table = Table(data, colWidths=[0.8 * cm, 5.1 * cm, 6.5 * cm, 2.4 * cm], repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONT", (0, 0), (-1, -1), "Times-Roman", 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)
    story.append(PageBreak())

    for head, body in [
        ("BAB VII HASIL TAMPILAN", "Screenshot APK Android release belum tersedia. Folder screenshot berisi catatan blocker dan perlu dilengkapi setelah Android SDK tersedia."),
        ("BAB VIII CARA MENJALANKAN", "Laravel dijalankan dengan php artisan serve --host=0.0.0.0 --port=8000. Flutter web memakai API_BASE_URL http://127.0.0.1:8000/api. Android emulator memakai http://10.0.2.2:8000/api."),
        ("BAB IX PENUTUP", "Fitur coding telah selesai dan validasi source lulus. Rilis Android masih terblokir Android SDK/device."),
        ("LAMPIRAN", "GitHub Flutter: [ISI LINK GITHUB FLUTTER SETELAH PUSH]. GitHub Laravel: [ISI LINK GITHUB LARAVEL SETELAH PUSH]."),
    ]:
        story.append(Paragraph(head, styles["Head"]))
        story.append(Paragraph(body, styles["BodyJustify"]))
        story.append(Spacer(1, 0.3 * cm))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Times-Roman", 9)
        canvas.drawCentredString(A4[0] / 2, 1.5 * cm, str(doc.page))
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(pdf),
        pagesize=A4,
        leftMargin=4 * cm,
        rightMargin=3 * cm,
        topMargin=3 * cm,
        bottomMargin=3 * cm,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return pdf


def write_checksums(paths: list[Path]) -> None:
    lines = ["CHECKSUM SHA-256 SMARTKASIR QR", ""]
    for path in paths:
        if path.exists():
            lines.append(f"{sha256(path)}  {path.name}")
    apk = APK_DIR / "SmartKasirQR-v1.0.0.apk"
    if not apk.exists():
        lines.append("APK belum tersedia karena Android SDK/device belum tersedia.")
    (SUBMISSION / "CHECKSUM_SHA256.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()
    write_text_files()
    docx = build_docx()
    pdf = build_pdf()
    write_checksums([docx, pdf])
    print(f"DOCX={docx}")
    print(f"PDF={pdf}")
    print(f"DOCX_SHA256={sha256(docx)}")
    print(f"PDF_SHA256={sha256(pdf)}")


if __name__ == "__main__":
    main()
